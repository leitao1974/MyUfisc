
from docx import Document
from io import BytesIO

def gerar_relatorio(dados):

    doc = Document()

    doc.add_heading("AUTO TÉCNICO DE FISCALIZAÇÃO AMBIENTAL", level=1)

    doc.add_heading("1. Identificação", level=2)
    doc.add_paragraph(f"Local: {dados['local']}")
    doc.add_paragraph(f"Coordenadas: {dados['lat']} / {dados['lon']}")

    doc.add_heading("2. Descrição dos Factos", level=2)
    doc.add_paragraph(dados["descricao"])

    doc.add_heading("3. Regimes Legais Aplicáveis", level=2)

    for r in dados["regimes"]:
        doc.add_paragraph(
            f"{r['regime']} — {r['diploma']} ({r['artigo']})"
        )

    doc.add_heading("4. Conclusão", level=2)
    doc.add_paragraph(
        "Da análise efetuada conclui-se existir indícios de infração aos regimes jurídicos identificados."
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
